#!/usr/bin/env python3
"""
══════════════════════════════════════════════════════════════════════════════════
                     SRM INTELLIGENT PROCESSOR v4.0
                  Procesador Universal Multi-Industria
══════════════════════════════════════════════════════════════════════════════════

DESCRIPCIÓN:
    Procesador universal de catálogos que transforma cualquier fuente de datos
    en información técnica estructurada lista para e-commerce.

FORMATOS SOPORTADOS:
    ✓ PDF      - Catálogos, fichas técnicas, listas de precios
    ✓ Excel    - .xlsx, .xls con datos tabulares
    ✓ CSV      - Archivos delimitados
    ✓ Word     - .docx con tablas y texto
    ✓ TXT      - Texto plano estructurado
    ✓ Imágenes - JPG, PNG, WEBP (análisis con Vision AI)
    ✓ ZIP      - Archivos comprimidos con múltiples fuentes
    ✓ URL      - Scraping de páginas web de catálogos

PIPELINE:
    1. INGESTA      → Recepción y detección automática de formato
    2. EXTRACCIÓN   → Parsing inteligente de datos
    3. NORMALIZACIÓN → Limpieza y estandarización
    4. UNIFICACIÓN  → Clasificación bajo taxonomía técnica
    5. ENRIQUECIMIENTO → Fitment, atributos técnicos, descripciones
    6. FICHA 360°   → Generación de ficha completa para publicación

INDUSTRIAS SOPORTADAS:
    - Autopartes (motos, carros, camiones)
    - Ferretería y construcción
    - Electrónica y tecnología
    - Hogar y decoración
    - Industrial y maquinaria
    - Genérico (configurable)

INSTALACIÓN:
    pip install openai pandas openpyxl python-docx beautifulsoup4 requests pillow opencv-python

USO:
    python3 srm_intelligent_processor.py <archivo_o_url> [opciones]

EJEMPLOS:
    python3 srm_intelligent_processor.py catalogo.pdf --industry autopartes
    python3 srm_intelligent_processor.py productos.xlsx --output /data/output
    python3 srm_intelligent_processor.py "https://ejemplo.com/catalogo" --scrape
    python3 srm_intelligent_processor.py archivos.zip --prefix PROV

AUTOR: SRM Team
VERSIÓN: 4.0
══════════════════════════════════════════════════════════════════════════════════
"""

import os
import sys
import gc
import re
import json
import csv
import base64
import time
import hashlib
import tempfile
import zipfile
import mimetypes
import urllib.parse
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple, Optional, Any, Union
from dataclasses import dataclass, field, asdict
from abc import ABC, abstractmethod
import subprocess
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed

# ============================================================================
# VERIFICACIÓN DE DEPENDENCIAS
# ============================================================================

DEPENDENCIES = {
    'pandas': 'pandas',
    'openpyxl': 'openpyxl',
    'cv2': 'opencv-python',
    'numpy': 'numpy',
    'PIL': 'pillow',
    'openai': 'openai',
    'docx': 'python-docx',
    'bs4': 'beautifulsoup4',
    'requests': 'requests',
}

def check_dependencies():
    """Verifica e importa dependencias."""
    missing = []

    for module, package in DEPENDENCIES.items():
        try:
            __import__(module)
        except ImportError:
            missing.append(package)

    if missing:
        print(f"❌ Dependencias faltantes:")
        print(f"   pip install {' '.join(missing)}")
        sys.exit(1)

check_dependencies()

import pandas as pd
import numpy as np
import cv2
from PIL import Image
from openai import OpenAI
from docx import Document
from bs4 import BeautifulSoup
import requests

# ============================================================================
# CONFIGURACIÓN GLOBAL
# ============================================================================

VERSION = "4.0"
SCRIPT_NAME = "SRM Intelligent Processor"

# Directorios
DEFAULT_OUTPUT_DIR = "/tmp/srm_output"
DEFAULT_TEMP_DIR = "/tmp/srm_temp"
DEFAULT_CACHE_DIR = "/tmp/srm_cache"

# API Configuration
VISION_MODEL = "gpt-4o"
TEXT_MODEL = "gpt-4o-mini"
MAX_TOKENS = 4096
MAX_RETRIES = 5
RETRY_DELAY = 2
REQUEST_TIMEOUT = 120

# Processing
BATCH_SIZE = 10
MAX_WORKERS = 4
CHECKPOINT_INTERVAL = 5

# Scraping
DEFAULT_USER_AGENT = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
SCRAPE_DELAY = 1.5
MAX_PAGES_SCRAPE = 50

# File limits
MAX_FILE_SIZE_MB = 100
MAX_IMAGES_PER_ZIP = 500


# ============================================================================
# TAXONOMÍA MULTI-INDUSTRIA
# ============================================================================

INDUSTRY_TAXONOMIES = {
    "autopartes": {
        "name": "Autopartes y Vehículos",
        "categories": {
            "motor": ["motor", "culata", "pistón", "biela", "cigüeñal", "válvula", "cilindro",
                     "empaque", "junta", "retén", "árbol", "leva", "cárter", "balancín"],
            "frenos": ["freno", "pastilla", "disco", "zapata", "caliper", "mordaza", "tambor",
                      "bomba freno", "líquido freno", "manguera freno"],
            "suspension": ["suspensión", "amortiguador", "resorte", "espiral", "buje",
                          "rótula", "terminal", "barra", "brazo"],
            "transmision": ["transmisión", "caja", "embrague", "clutch", "cadena", "piñón",
                           "catalina", "correa", "variador", "eje", "cruceta"],
            "electrico": ["eléctrico", "batería", "alternador", "motor arranque", "bobina",
                         "bujía", "cable", "faro", "luz", "bombillo", "fusible", "relay"],
            "carroceria": ["carrocería", "guardafango", "parachoques", "capó", "puerta",
                          "espejo", "vidrio", "moldura", "manija"],
            "direccion": ["dirección", "volante", "columna", "cremallera", "bomba dirección"],
            "refrigeracion": ["refrigeración", "radiador", "termostato", "bomba agua",
                             "ventilador", "manguera", "anticongelante"],
            "escape": ["escape", "silenciador", "catalizador", "tubo", "colector"],
            "accesorios": ["accesorio", "tapete", "forro", "portavasos", "organizador"],
        },
        "fitment_fields": ["marca", "modelo", "año", "motor", "cilindraje", "posicion"],
    },

    "ferreteria": {
        "name": "Ferretería y Construcción",
        "categories": {
            "herramientas_manuales": ["martillo", "destornillador", "llave", "alicate", "pinza",
                                      "serrucho", "cincel", "lima", "nivel", "metro", "flexómetro"],
            "herramientas_electricas": ["taladro", "esmeril", "sierra", "lijadora", "pulidora",
                                        "rotomartillo", "caladora", "fresadora"],
            "fijacion": ["tornillo", "clavo", "taco", "ancla", "remache", "grapa", "tuerca",
                        "arandela", "perno", "abrazadera"],
            "plomeria": ["tubo", "codo", "tee", "unión", "válvula", "llave paso", "grifo",
                        "sifón", "sanitario", "lavamanos"],
            "electricos": ["cable", "interruptor", "tomacorriente", "breaker", "caja",
                          "canaleta", "tubo conduit", "cinta", "bombillo"],
            "pinturas": ["pintura", "esmalte", "vinilo", "barniz", "sellador", "brocha",
                        "rodillo", "espátula", "masilla", "lija"],
            "construccion": ["cemento", "arena", "grava", "ladrillo", "bloque", "varilla",
                            "malla", "alambre", "impermeabilizante"],
            "cerrajeria": ["cerradura", "candado", "chapa", "llave", "bisagra", "pasador"],
            "jardineria": ["manguera", "aspersor", "tijera", "pala", "rastrillo", "machete"],
            "seguridad": ["casco", "guante", "gafa", "bota", "arnés", "chaleco", "tapaoídos"],
        },
        "fitment_fields": ["material", "medida", "tipo", "uso", "norma"],
    },

    "electronica": {
        "name": "Electrónica y Tecnología",
        "categories": {
            "computacion": ["computador", "laptop", "monitor", "teclado", "mouse", "disco",
                           "memoria", "procesador", "tarjeta", "fuente"],
            "celulares": ["celular", "smartphone", "tablet", "cargador", "cable", "forro",
                         "mica", "batería", "auricular", "parlante"],
            "audio_video": ["televisor", "teatro", "parlante", "amplificador", "micrófono",
                           "cámara", "proyector", "consola"],
            "redes": ["router", "switch", "cable red", "access point", "modem", "rack"],
            "componentes": ["resistencia", "capacitor", "transistor", "diodo", "led",
                           "circuito", "sensor", "relay", "transformador"],
            "electrodomesticos": ["nevera", "lavadora", "horno", "microondas", "licuadora",
                                 "ventilador", "aire acondicionado", "plancha"],
            "iluminacion": ["bombillo", "lámpara", "reflector", "cinta led", "driver"],
            "energia": ["batería", "ups", "inversor", "panel solar", "regulador"],
            "cables": ["cable", "conector", "adaptador", "extensión", "regleta", "enchufe"],
            "accesorios": ["estuche", "soporte", "base", "organizador", "limpiador"],
        },
        "fitment_fields": ["marca", "modelo", "compatibilidad", "voltaje", "potencia"],
    },

    "hogar": {
        "name": "Hogar y Decoración",
        "categories": {
            "muebles": ["sofá", "silla", "mesa", "cama", "armario", "escritorio", "estante"],
            "cocina": ["olla", "sartén", "plato", "vaso", "cubierto", "tabla", "colador"],
            "baño": ["toalla", "cortina", "tapete", "dispensador", "espejo", "organizador"],
            "decoracion": ["cuadro", "jarrón", "cojín", "lámpara", "reloj", "espejo"],
            "textiles": ["sábana", "cobija", "almohada", "cortina", "mantel", "tapete"],
            "almacenamiento": ["caja", "canasta", "organizador", "gancho", "perchero"],
            "jardin": ["maceta", "manguera", "silla exterior", "sombrilla", "parrilla"],
            "limpieza": ["escoba", "trapeador", "balde", "cepillo", "esponja", "detergente"],
        },
        "fitment_fields": ["material", "color", "medidas", "estilo", "ambiente"],
    },

    "industrial": {
        "name": "Industrial y Maquinaria",
        "categories": {
            "motores": ["motor eléctrico", "motor diesel", "motor gasolina", "reductor"],
            "bombas": ["bomba centrífuga", "bomba sumergible", "bomba neumática", "compresor"],
            "valvulas": ["válvula", "electroválvula", "regulador", "manómetro", "filtro"],
            "transmision": ["polea", "correa", "cadena", "sprocket", "chumacera", "rodamiento"],
            "neumatica": ["cilindro", "válvula neumática", "racor", "manguera", "FRL"],
            "hidraulica": ["cilindro hidráulico", "bomba hidráulica", "válvula", "aceite"],
            "automatizacion": ["PLC", "sensor", "actuador", "variador", "HMI", "contactor"],
            "soldadura": ["soldadora", "electrodo", "alambre", "gas", "careta", "guante"],
            "seguridad": ["EPP", "señalización", "extintor", "botiquín", "camilla"],
            "mantenimiento": ["lubricante", "grasa", "limpiador", "sellante", "adhesivo"],
        },
        "fitment_fields": ["marca", "modelo", "potencia", "voltaje", "capacidad", "presion"],
    },

    "generico": {
        "name": "Catálogo General",
        "categories": {
            "productos": ["producto", "artículo", "item", "referencia"],
        },
        "fitment_fields": ["marca", "modelo", "tipo", "caracteristica"],
    },
}


# ============================================================================
# LOGGING
# ============================================================================

class Colors:
    HEADER = '\033[95m'
    BLUE = '\033[94m'
    CYAN = '\033[96m'
    GREEN = '\033[92m'
    YELLOW = '\033[93m'
    RED = '\033[91m'
    BOLD = '\033[1m'
    DIM = '\033[2m'
    RESET = '\033[0m'


class Logger:
    LEVELS = {
        'debug': (Colors.DIM, ''),
        'info': (Colors.CYAN, ''),
        'success': (Colors.GREEN, '✓'),
        'warning': (Colors.YELLOW, '⚠'),
        'error': (Colors.RED, '✗'),
        'header': (Colors.BOLD, ''),
        'step': (Colors.BLUE, '→'),
    }

    def __init__(self, verbose: bool = True):
        self.verbose = verbose
        self.start_time = time.time()

    def log(self, msg: str, level: str = 'info'):
        if not self.verbose and level == 'debug':
            return
        color, icon = self.LEVELS.get(level, (Colors.RESET, ''))
        ts = datetime.now().strftime('%H:%M:%S')
        prefix = f"{icon} " if icon else ""
        print(f"{color}[{ts}] {prefix}{msg}{Colors.RESET}", flush=True)

    def step(self, step_num: int, total: int, name: str):
        print(f"\n{Colors.BOLD}{'='*60}")
        print(f"  PASO {step_num}/{total}: {name.upper()}")
        print(f"{'='*60}{Colors.RESET}\n")

    def elapsed(self) -> str:
        elapsed = time.time() - self.start_time
        if elapsed < 60:
            return f"{elapsed:.1f}s"
        return f"{elapsed/60:.1f}m"


log = Logger()


# ============================================================================
# UTILIDADES
# ============================================================================

def ensure_dir(path: str) -> str:
    os.makedirs(path, exist_ok=True)
    return path


def file_hash(filepath: str, length: int = 12) -> str:
    hasher = hashlib.md5()
    with open(filepath, 'rb') as f:
        for chunk in iter(lambda: f.read(65536), b''):
            hasher.update(chunk)
    return hasher.hexdigest()[:length]


def detect_file_type(filepath: str) -> str:
    """Detecta el tipo de archivo."""
    ext = Path(filepath).suffix.lower()

    type_map = {
        '.pdf': 'pdf',
        '.xlsx': 'excel',
        '.xls': 'excel',
        '.csv': 'csv',
        '.docx': 'word',
        '.doc': 'word',
        '.txt': 'txt',
        '.jpg': 'image',
        '.jpeg': 'image',
        '.png': 'image',
        '.webp': 'image',
        '.gif': 'image',
        '.zip': 'zip',
        '.rar': 'zip',
        '.7z': 'zip',
    }

    return type_map.get(ext, 'unknown')


def clean_text(text: Any) -> str:
    if not isinstance(text, str):
        text = str(text) if text is not None else ""
    text = text.strip()
    text = re.sub(r'\s+', ' ', text)
    return text


def clean_price(value: Any) -> float:
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        cleaned = re.sub(r'[^\d.,]', '', value)
        cleaned = cleaned.replace(',', '.')
        if cleaned.count('.') > 1:
            parts = cleaned.split('.')
            cleaned = ''.join(parts[:-1]) + '.' + parts[-1]
        try:
            return float(cleaned) if cleaned else 0.0
        except ValueError:
            return 0.0
    return 0.0


def normalize_category(text: str, industry: str = "generico") -> str:
    """Normaliza categoría según industria."""
    if not text:
        return "OTROS"

    text_lower = text.lower()
    taxonomy = INDUSTRY_TAXONOMIES.get(industry, INDUSTRY_TAXONOMIES["generico"])

    for category, keywords in taxonomy["categories"].items():
        for keyword in keywords:
            if keyword in text_lower:
                return category.upper().replace("_", " ")

    return "OTROS"


# ============================================================================
# MODELOS DE DATOS
# ============================================================================

@dataclass
class ProductData:
    """Datos de producto extraído."""
    id: str = ""
    codigo: str = ""
    sku: str = ""
    nombre: str = ""
    descripcion: str = ""
    descripcion_corta: str = ""
    precio: float = 0.0
    precio_anterior: float = 0.0
    categoria: str = ""
    subcategoria: str = ""
    marca: str = ""
    modelo: str = ""
    imagen: str = ""
    imagenes: List[str] = field(default_factory=list)
    atributos: Dict[str, Any] = field(default_factory=dict)
    fitment: Dict[str, Any] = field(default_factory=dict)
    tags: List[str] = field(default_factory=list)
    stock: int = 0
    peso: float = 0.0
    dimensiones: str = ""
    fuente: str = ""
    pagina: int = 0
    confianza: float = 0.0

    def to_dict(self) -> dict:
        d = asdict(self)
        d['imagenes'] = ','.join(self.imagenes) if self.imagenes else ''
        d['tags'] = ','.join(self.tags) if self.tags else ''
        d['atributos'] = json.dumps(self.atributos, ensure_ascii=False) if self.atributos else ''
        d['fitment'] = json.dumps(self.fitment, ensure_ascii=False) if self.fitment else ''
        return d

    def is_valid(self) -> bool:
        return bool(self.codigo or self.nombre)


@dataclass
class ProcessingResult:
    """Resultado del procesamiento."""
    success: bool = False
    products: List[ProductData] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)
    source_file: str = ""
    source_type: str = ""
    processing_time: float = 0.0


# ============================================================================
# PROCESADORES DE ARCHIVOS
# ============================================================================

class BaseProcessor(ABC):
    """Clase base para procesadores de archivos."""

    def __init__(self, config: dict):
        self.config = config
        self.industry = config.get('industry', 'generico')
        self.prefix = config.get('prefix', 'SRM')
        self.output_dir = config.get('output_dir', DEFAULT_OUTPUT_DIR)

    @abstractmethod
    def process(self, filepath: str) -> ProcessingResult:
        pass

    def _get_openai_client(self) -> OpenAI:
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("OPENAI_API_KEY no configurada")
        return OpenAI(api_key=api_key)


class PDFProcessor(BaseProcessor):
    """Procesa archivos PDF."""

    def process(self, filepath: str) -> ProcessingResult:
        log.log(f"Procesando PDF: {Path(filepath).name}", "step")
        result = ProcessingResult(source_file=filepath, source_type="pdf")
        start_time = time.time()

        try:
            client = self._get_openai_client()
            pages_dir = ensure_dir(os.path.join(self.output_dir, "pages"))
            crops_dir = ensure_dir(os.path.join(self.output_dir, "crops"))

            # Obtener número de páginas
            page_count = self._get_page_count(filepath)
            pages_to_process = self.config.get('pages', list(range(1, (page_count or 50) + 1)))

            log.log(f"   Páginas a procesar: {len(pages_to_process)}")

            for i, page_num in enumerate(pages_to_process, 1):
                log.log(f"   [{i}/{len(pages_to_process)}] Página {page_num}...")

                # Convertir página
                page_img = os.path.join(pages_dir, f"page_{page_num:03d}.jpg")
                if not self._convert_page(filepath, page_num, page_img):
                    continue

                # Extraer productos con Vision
                products = self._extract_with_vision(client, page_img, page_num)

                # Detectar y guardar crops
                crops = self._detect_and_crop(page_img, crops_dir, page_num)

                # Asociar crops a productos
                if crops and products:
                    products = self._associate_crops(products, crops)

                result.products.extend(products)
                log.log(f"      {len(products)} productos extraídos", "success")

                gc.collect()

            result.success = True

        except Exception as e:
            result.errors.append(str(e))
            log.log(f"Error procesando PDF: {e}", "error")

        result.processing_time = time.time() - start_time
        return result

    def _get_page_count(self, filepath: str) -> Optional[int]:
        try:
            result = subprocess.run(['pdfinfo', filepath], capture_output=True, text=True, timeout=30)
            for line in result.stdout.split('\n'):
                if line.startswith('Pages:'):
                    return int(line.split(':')[1].strip())
        except:
            pass
        return None

    def _convert_page(self, pdf_path: str, page_num: int, output_path: str) -> bool:
        try:
            base = output_path.rsplit('.', 1)[0]
            dpi = self.config.get('dpi', 150)
            cmd = ['pdftoppm', '-jpeg', '-r', str(dpi), '-f', str(page_num),
                   '-l', str(page_num), '-singlefile', pdf_path, base]
            result = subprocess.run(cmd, capture_output=True, timeout=120)
            return result.returncode == 0 and os.path.exists(output_path)
        except Exception as e:
            log.log(f"Error convirtiendo página: {e}", "warning")
            return False

    def _extract_with_vision(self, client: OpenAI, image_path: str, page_num: int) -> List[ProductData]:
        prompt = self._get_extraction_prompt()

        with open(image_path, 'rb') as f:
            image_b64 = base64.b64encode(f.read()).decode('utf-8')

        for attempt in range(MAX_RETRIES):
            try:
                response = client.chat.completions.create(
                    model=VISION_MODEL,
                    messages=[{
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {"type": "image_url", "image_url": {
                                "url": f"data:image/jpeg;base64,{image_b64}",
                                "detail": "high"
                            }}
                        ]
                    }],
                    max_tokens=MAX_TOKENS,
                    response_format={"type": "json_object"},
                    timeout=REQUEST_TIMEOUT
                )

                data = json.loads(response.choices[0].message.content)
                products = []

                for p in data.get("productos", []):
                    product = ProductData(
                        codigo=clean_text(p.get('codigo', '')),
                        nombre=clean_text(p.get('nombre', '')),
                        descripcion=clean_text(p.get('descripcion', '')),
                        precio=clean_price(p.get('precio', 0)),
                        categoria=normalize_category(p.get('categoria', ''), self.industry),
                        marca=clean_text(p.get('marca', '')),
                        pagina=page_num,
                        fuente=Path(image_path).name
                    )

                    if p.get('posicion_vertical'):
                        product.atributos['posicion_y'] = int(p['posicion_vertical'])

                    if product.is_valid():
                        product.sku = f"{self.prefix}-{product.codigo}" if product.codigo else ""
                        products.append(product)

                return products

            except Exception as e:
                if 'rate_limit' in str(e).lower():
                    time.sleep(RETRY_DELAY * (2 ** attempt))
                else:
                    log.log(f"Error Vision API: {str(e)[:50]}", "warning")
                    time.sleep(RETRY_DELAY)

        return []

    def _get_extraction_prompt(self) -> str:
        industry_name = INDUSTRY_TAXONOMIES.get(self.industry, {}).get('name', 'productos')

        return f"""Analiza esta página de catálogo de {industry_name}.

EXTRAE todos los productos visibles con:
- codigo: Código/referencia del producto
- nombre: Nombre comercial
- descripcion: Especificaciones técnicas
- precio: Valor numérico sin símbolos
- categoria: Categoría del producto
- marca: Marca si es visible
- posicion_vertical: Posición 1-10 (1=arriba, 10=abajo)

RESPONDE JSON: {{"productos": [...]}}"""

    def _detect_and_crop(self, image_path: str, output_dir: str, page_num: int) -> List[dict]:
        try:
            img = cv2.imread(image_path)
            if img is None:
                return []

            height, width = img.shape[:2]
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            blurred = cv2.GaussianBlur(gray, (5, 5), 0)
            edges = cv2.Canny(blurred, 30, 100)
            kernel = np.ones((5, 5), np.uint8)
            dilated = cv2.dilate(edges, kernel, iterations=2)
            contours, _ = cv2.findContours(dilated, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

            crops = []
            for i, contour in enumerate(contours):
                x, y, w, h = cv2.boundingRect(contour)
                area = w * h

                if w < 50 or h < 50 or area > (height * width * 0.5):
                    continue

                crop_img = img[y:y+h, x:x+w]
                filename = f"{self.prefix}_p{page_num:03d}_c{i+1:02d}.jpg"
                path = os.path.join(output_dir, filename)
                cv2.imwrite(path, crop_img, [cv2.IMWRITE_JPEG_QUALITY, 90])

                crops.append({
                    'filename': filename,
                    'y_normalized': y / height,
                    'assigned': False
                })

            return crops[:20]
        except Exception as e:
            log.log(f"Error detectando crops: {e}", "warning")
            return []

    def _associate_crops(self, products: List[ProductData], crops: List[dict]) -> List[ProductData]:
        for product in sorted(products, key=lambda p: p.atributos.get('posicion_y', 5)):
            pos_y = (product.atributos.get('posicion_y', 5) - 1) / 9

            best_crop = None
            best_distance = float('inf')

            for crop in crops:
                if crop['assigned']:
                    continue
                distance = abs(pos_y - crop['y_normalized'])
                if distance < best_distance:
                    best_distance = distance
                    best_crop = crop

            if best_crop and best_distance < 0.25:
                product.imagen = best_crop['filename']
                best_crop['assigned'] = True

        return products


class ExcelProcessor(BaseProcessor):
    """Procesa archivos Excel."""

    def process(self, filepath: str) -> ProcessingResult:
        log.log(f"Procesando Excel: {Path(filepath).name}", "step")
        result = ProcessingResult(source_file=filepath, source_type="excel")
        start_time = time.time()

        try:
            # Leer todas las hojas
            xls = pd.ExcelFile(filepath)

            for sheet_name in xls.sheet_names:
                log.log(f"   Hoja: {sheet_name}")
                df = pd.read_excel(filepath, sheet_name=sheet_name)

                if df.empty:
                    continue

                # Mapear columnas
                products = self._extract_from_dataframe(df, sheet_name)
                result.products.extend(products)
                log.log(f"      {len(products)} productos", "success")

            result.success = True

        except Exception as e:
            result.errors.append(str(e))
            log.log(f"Error procesando Excel: {e}", "error")

        result.processing_time = time.time() - start_time
        return result

    def _extract_from_dataframe(self, df: pd.DataFrame, source: str) -> List[ProductData]:
        # Normalizar nombres de columnas
        df.columns = [str(c).lower().strip() for c in df.columns]

        # Mapeo de columnas comunes
        column_maps = {
            'codigo': ['codigo', 'code', 'ref', 'referencia', 'sku', 'id', 'part_number'],
            'nombre': ['nombre', 'name', 'descripcion', 'description', 'producto', 'product', 'titulo'],
            'precio': ['precio', 'price', 'valor', 'value', 'costo', 'cost', 'pvp'],
            'categoria': ['categoria', 'category', 'familia', 'family', 'tipo', 'type', 'linea'],
            'marca': ['marca', 'brand', 'fabricante', 'manufacturer'],
            'stock': ['stock', 'cantidad', 'qty', 'inventory', 'existencia'],
        }

        mapped = {}
        for field, options in column_maps.items():
            for opt in options:
                if opt in df.columns:
                    mapped[field] = opt
                    break

        products = []
        for _, row in df.iterrows():
            product = ProductData(
                codigo=clean_text(row.get(mapped.get('codigo', ''), '')),
                nombre=clean_text(row.get(mapped.get('nombre', ''), '')),
                precio=clean_price(row.get(mapped.get('precio', ''), 0)),
                categoria=normalize_category(
                    str(row.get(mapped.get('categoria', ''), '')),
                    self.industry
                ),
                marca=clean_text(row.get(mapped.get('marca', ''), '')),
                stock=int(row.get(mapped.get('stock', ''), 0) or 0),
                fuente=source
            )

            if product.is_valid():
                product.sku = f"{self.prefix}-{product.codigo}" if product.codigo else ""
                products.append(product)

        return products


class CSVProcessor(BaseProcessor):
    """Procesa archivos CSV."""

    def process(self, filepath: str) -> ProcessingResult:
        log.log(f"Procesando CSV: {Path(filepath).name}", "step")
        result = ProcessingResult(source_file=filepath, source_type="csv")
        start_time = time.time()

        try:
            # Detectar separador
            with open(filepath, 'r', encoding='utf-8-sig') as f:
                sample = f.read(4096)

            separators = [';', ',', '\t', '|']
            sep = max(separators, key=lambda s: sample.count(s))

            df = pd.read_csv(filepath, sep=sep, encoding='utf-8-sig')

            # Usar mismo procesamiento que Excel
            excel_proc = ExcelProcessor(self.config)
            products = excel_proc._extract_from_dataframe(df, Path(filepath).name)

            result.products = products
            result.success = True
            log.log(f"   {len(products)} productos extraídos", "success")

        except Exception as e:
            result.errors.append(str(e))
            log.log(f"Error procesando CSV: {e}", "error")

        result.processing_time = time.time() - start_time
        return result


class WordProcessor(BaseProcessor):
    """Procesa archivos Word."""

    def process(self, filepath: str) -> ProcessingResult:
        log.log(f"Procesando Word: {Path(filepath).name}", "step")
        result = ProcessingResult(source_file=filepath, source_type="word")
        start_time = time.time()

        try:
            doc = Document(filepath)

            # Extraer tablas
            for table in doc.tables:
                products = self._extract_from_table(table)
                result.products.extend(products)

            # Si no hay tablas, procesar texto
            if not result.products:
                text = '\n'.join([p.text for p in doc.paragraphs])
                products = self._extract_from_text_with_ai(text)
                result.products.extend(products)

            result.success = True
            log.log(f"   {len(result.products)} productos extraídos", "success")

        except Exception as e:
            result.errors.append(str(e))
            log.log(f"Error procesando Word: {e}", "error")

        result.processing_time = time.time() - start_time
        return result

    def _extract_from_table(self, table) -> List[ProductData]:
        products = []
        headers = []

        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]

            if i == 0:
                headers = [c.lower() for c in cells]
                continue

            if len(cells) >= 2:
                product = ProductData(
                    codigo=cells[0] if len(cells) > 0 else '',
                    nombre=cells[1] if len(cells) > 1 else '',
                    precio=clean_price(cells[2]) if len(cells) > 2 else 0,
                    fuente="word_table"
                )
                if product.is_valid():
                    product.sku = f"{self.prefix}-{product.codigo}" if product.codigo else ""
                    products.append(product)

        return products

    def _extract_from_text_with_ai(self, text: str) -> List[ProductData]:
        if len(text) < 50:
            return []

        try:
            client = self._get_openai_client()

            response = client.chat.completions.create(
                model=TEXT_MODEL,
                messages=[{
                    "role": "user",
                    "content": f"""Extrae productos del siguiente texto.

RESPONDE JSON: {{"productos": [{{"codigo": "", "nombre": "", "descripcion": "", "precio": 0}}]}}

TEXTO:
{text[:8000]}"""
                }],
                response_format={"type": "json_object"},
                max_tokens=2000
            )

            data = json.loads(response.choices[0].message.content)
            products = []

            for p in data.get("productos", []):
                product = ProductData(
                    codigo=clean_text(p.get('codigo', '')),
                    nombre=clean_text(p.get('nombre', '')),
                    descripcion=clean_text(p.get('descripcion', '')),
                    precio=clean_price(p.get('precio', 0)),
                    fuente="word_text"
                )
                if product.is_valid():
                    product.sku = f"{self.prefix}-{product.codigo}" if product.codigo else ""
                    products.append(product)

            return products

        except Exception as e:
            log.log(f"Error extrayendo de texto: {e}", "warning")
            return []


class TXTProcessor(BaseProcessor):
    """Procesa archivos de texto."""

    def process(self, filepath: str) -> ProcessingResult:
        log.log(f"Procesando TXT: {Path(filepath).name}", "step")
        result = ProcessingResult(source_file=filepath, source_type="txt")
        start_time = time.time()

        try:
            with open(filepath, 'r', encoding='utf-8-sig') as f:
                text = f.read()

            # Usar el procesador de Word para extraer con AI
            word_proc = WordProcessor(self.config)
            products = word_proc._extract_from_text_with_ai(text)

            result.products = products
            result.success = True
            log.log(f"   {len(products)} productos extraídos", "success")

        except Exception as e:
            result.errors.append(str(e))
            log.log(f"Error procesando TXT: {e}", "error")

        result.processing_time = time.time() - start_time
        return result


class ImageProcessor(BaseProcessor):
    """Procesa imágenes individuales."""

    def process(self, filepath: str) -> ProcessingResult:
        log.log(f"Procesando imagen: {Path(filepath).name}", "step")
        result = ProcessingResult(source_file=filepath, source_type="image")
        start_time = time.time()

        try:
            client = self._get_openai_client()

            with open(filepath, 'rb') as f:
                image_b64 = base64.b64encode(f.read()).decode('utf-8')

            response = client.chat.completions.create(
                model=VISION_MODEL,
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": self._get_image_prompt()},
                        {"type": "image_url", "image_url": {
                            "url": f"data:image/jpeg;base64,{image_b64}",
                            "detail": "high"
                        }}
                    ]
                }],
                max_tokens=2000,
                response_format={"type": "json_object"}
            )

            data = json.loads(response.choices[0].message.content)

            product = ProductData(
                codigo=clean_text(data.get('codigo', '')),
                nombre=clean_text(data.get('nombre', '')),
                descripcion=clean_text(data.get('descripcion', '')),
                categoria=normalize_category(data.get('categoria', ''), self.industry),
                marca=clean_text(data.get('marca', '')),
                imagen=Path(filepath).name,
                fuente=Path(filepath).name,
                atributos=data.get('atributos', {}),
                tags=data.get('tags', [])
            )

            if product.nombre:
                product.sku = f"{self.prefix}-{product.codigo}" if product.codigo else f"{self.prefix}-IMG"
                result.products.append(product)

            result.success = True
            log.log(f"   Producto identificado: {product.nombre[:50]}", "success")

        except Exception as e:
            result.errors.append(str(e))
            log.log(f"Error procesando imagen: {e}", "error")

        result.processing_time = time.time() - start_time
        return result

    def _get_image_prompt(self) -> str:
        industry_name = INDUSTRY_TAXONOMIES.get(self.industry, {}).get('name', 'productos')

        return f"""Analiza esta imagen de un producto de {industry_name}.

Identifica:
- codigo: Código si es visible (en etiqueta, grabado, etc.)
- nombre: Nombre comercial descriptivo para catálogo
- descripcion: Descripción técnica detallada
- categoria: Categoría del producto
- marca: Marca si es identificable
- atributos: {{material, color, medidas, etc.}}
- tags: [palabras clave para búsqueda]

RESPONDE JSON con estos campos."""


class ZIPProcessor(BaseProcessor):
    """Procesa archivos ZIP."""

    def process(self, filepath: str) -> ProcessingResult:
        log.log(f"Procesando ZIP: {Path(filepath).name}", "step")
        result = ProcessingResult(source_file=filepath, source_type="zip")
        start_time = time.time()

        try:
            extract_dir = ensure_dir(os.path.join(self.output_dir, "extracted"))

            with zipfile.ZipFile(filepath, 'r') as zip_ref:
                zip_ref.extractall(extract_dir)

            # Procesar cada archivo extraído
            for root, dirs, files in os.walk(extract_dir):
                for file in files[:MAX_IMAGES_PER_ZIP]:
                    file_path = os.path.join(root, file)
                    file_type = detect_file_type(file_path)

                    if file_type == 'unknown':
                        continue

                    processor = get_processor(file_type, self.config)
                    if processor:
                        sub_result = processor.process(file_path)
                        result.products.extend(sub_result.products)
                        result.errors.extend(sub_result.errors)

            result.success = True
            log.log(f"   Total productos del ZIP: {len(result.products)}", "success")

        except Exception as e:
            result.errors.append(str(e))
            log.log(f"Error procesando ZIP: {e}", "error")

        result.processing_time = time.time() - start_time
        return result


class URLProcessor(BaseProcessor):
    """Procesa URLs (scraping)."""

    def process(self, url: str) -> ProcessingResult:
        log.log(f"Scraping URL: {url}", "step")
        result = ProcessingResult(source_file=url, source_type="url")
        start_time = time.time()

        try:
            headers = {'User-Agent': DEFAULT_USER_AGENT}
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()

            soup = BeautifulSoup(response.text, 'html.parser')

            # Intentar detectar productos en la página
            products = self._extract_products_from_html(soup, url)

            if not products:
                # Usar AI para extraer
                products = self._extract_with_ai(response.text, url)

            result.products = products
            result.success = True
            log.log(f"   {len(products)} productos encontrados", "success")

            # Buscar enlaces a más páginas
            if self.config.get('scrape_pagination', False):
                next_urls = self._find_pagination_links(soup, url)
                for next_url in next_urls[:MAX_PAGES_SCRAPE]:
                    time.sleep(SCRAPE_DELAY)
                    sub_result = self.process(next_url)
                    result.products.extend(sub_result.products)

        except Exception as e:
            result.errors.append(str(e))
            log.log(f"Error scraping URL: {e}", "error")

        result.processing_time = time.time() - start_time
        return result

    def _extract_products_from_html(self, soup: BeautifulSoup, url: str) -> List[ProductData]:
        products = []

        # Buscar patrones comunes de productos
        product_selectors = [
            '.product', '.producto', '.item', '.card',
            '[data-product]', '[itemtype*="Product"]',
            '.product-item', '.product-card'
        ]

        for selector in product_selectors:
            elements = soup.select(selector)
            if elements:
                for elem in elements:
                    product = self._parse_product_element(elem, url)
                    if product and product.is_valid():
                        products.append(product)
                break

        return products

    def _parse_product_element(self, elem, url: str) -> Optional[ProductData]:
        try:
            # Buscar nombre
            name_elem = elem.select_one('h1, h2, h3, h4, .title, .name, .product-name, .product-title')
            name = name_elem.get_text(strip=True) if name_elem else ''

            # Buscar precio
            price_elem = elem.select_one('.price, .precio, [class*="price"]')
            price_text = price_elem.get_text(strip=True) if price_elem else ''
            price = clean_price(price_text)

            # Buscar código/SKU
            sku_elem = elem.select_one('.sku, .codigo, [class*="sku"], [class*="ref"]')
            sku = sku_elem.get_text(strip=True) if sku_elem else ''

            # Buscar imagen
            img_elem = elem.select_one('img')
            img_src = img_elem.get('src', '') if img_elem else ''
            if img_src and not img_src.startswith('http'):
                img_src = urllib.parse.urljoin(url, img_src)

            if name:
                return ProductData(
                    codigo=sku,
                    nombre=name,
                    precio=price,
                    imagen=img_src,
                    fuente=url
                )
        except:
            pass

        return None

    def _extract_with_ai(self, html: str, url: str) -> List[ProductData]:
        try:
            client = self._get_openai_client()

            # Limpiar HTML
            soup = BeautifulSoup(html, 'html.parser')
            for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
                tag.decompose()
            text = soup.get_text(separator='\n', strip=True)[:10000]

            response = client.chat.completions.create(
                model=TEXT_MODEL,
                messages=[{
                    "role": "user",
                    "content": f"""Extrae productos de este contenido web.

RESPONDE JSON: {{"productos": [{{"codigo": "", "nombre": "", "descripcion": "", "precio": 0}}]}}

CONTENIDO:
{text}"""
                }],
                response_format={"type": "json_object"},
                max_tokens=2000
            )

            data = json.loads(response.choices[0].message.content)
            products = []

            for p in data.get("productos", []):
                product = ProductData(
                    codigo=clean_text(p.get('codigo', '')),
                    nombre=clean_text(p.get('nombre', '')),
                    descripcion=clean_text(p.get('descripcion', '')),
                    precio=clean_price(p.get('precio', 0)),
                    fuente=url
                )
                if product.is_valid():
                    product.sku = f"{self.prefix}-{product.codigo}" if product.codigo else ""
                    products.append(product)

            return products

        except Exception as e:
            log.log(f"Error AI scraping: {e}", "warning")
            return []

    def _find_pagination_links(self, soup: BeautifulSoup, base_url: str) -> List[str]:
        links = []
        pagination = soup.select('.pagination a, .pager a, [class*="page"] a')

        for link in pagination:
            href = link.get('href', '')
            if href and href not in ['#', '']:
                full_url = urllib.parse.urljoin(base_url, href)
                if full_url not in links:
                    links.append(full_url)

        return links


# ============================================================================
# FACTORY DE PROCESADORES
# ============================================================================

def get_processor(file_type: str, config: dict) -> Optional[BaseProcessor]:
    """Retorna el procesador apropiado según el tipo de archivo."""
    processors = {
        'pdf': PDFProcessor,
        'excel': ExcelProcessor,
        'csv': CSVProcessor,
        'word': WordProcessor,
        'txt': TXTProcessor,
        'image': ImageProcessor,
        'zip': ZIPProcessor,
        'url': URLProcessor,
    }

    processor_class = processors.get(file_type)
    if processor_class:
        return processor_class(config)
    return None


# ============================================================================
# PIPELINE DE PROCESAMIENTO
# ============================================================================

class SRMPipeline:
    """Pipeline completo de procesamiento SRM."""

    def __init__(self, config: dict):
        self.config = config
        self.output_dir = ensure_dir(config.get('output_dir', DEFAULT_OUTPUT_DIR))
        self.prefix = config.get('prefix', 'SRM')
        self.industry = config.get('industry', 'generico')
        self.all_products: List[ProductData] = []

    def process(self, source: str) -> Tuple[str, str]:
        """Ejecuta el pipeline completo."""
        self._print_banner()

        # 1. INGESTA
        log.step(1, 6, "INGESTA")
        result = self._ingest(source)

        if not result.success:
            log.log(f"Error en ingesta: {result.errors}", "error")
            return "", ""

        self.all_products = result.products
        log.log(f"Productos ingestados: {len(self.all_products)}", "success")

        # 2. EXTRACCIÓN (ya hecha en ingesta para la mayoría)
        log.step(2, 6, "EXTRACCIÓN")
        log.log(f"Datos extraídos de {result.source_type}", "success")

        # 3. NORMALIZACIÓN
        log.step(3, 6, "NORMALIZACIÓN")
        self._normalize()

        # 4. UNIFICACIÓN
        log.step(4, 6, "UNIFICACIÓN")
        self._unify()

        # 5. ENRIQUECIMIENTO
        log.step(5, 6, "ENRIQUECIMIENTO")
        self._enrich()

        # 6. FICHA 360°
        log.step(6, 6, "FICHA 360°")
        csv_path, json_path = self._export()

        self._print_summary(csv_path, json_path)

        return csv_path, json_path

    def _ingest(self, source: str) -> ProcessingResult:
        """Ingesta: detecta tipo y procesa."""
        # Detectar si es URL
        if source.startswith('http://') or source.startswith('https://'):
            processor = URLProcessor(self.config)
            return processor.process(source)

        # Detectar tipo de archivo
        if not os.path.exists(source):
            return ProcessingResult(errors=[f"Archivo no encontrado: {source}"])

        file_type = detect_file_type(source)
        log.log(f"Tipo detectado: {file_type}")

        processor = get_processor(file_type, self.config)
        if not processor:
            return ProcessingResult(errors=[f"Tipo no soportado: {file_type}"])

        return processor.process(source)

    def _normalize(self):
        """Normalización: limpieza y estandarización."""
        seen_codes = set()
        normalized = []

        for product in self.all_products:
            # Limpiar campos
            product.nombre = clean_text(product.nombre)
            product.descripcion = clean_text(product.descripcion)
            product.codigo = re.sub(r'[^a-zA-Z0-9-]', '', product.codigo)

            # Deduplicar por código
            if product.codigo:
                if product.codigo in seen_codes:
                    continue
                seen_codes.add(product.codigo)

            # Generar SKU si falta
            if not product.sku and product.codigo:
                product.sku = f"{self.prefix}-{product.codigo}"

            normalized.append(product)

        duplicates = len(self.all_products) - len(normalized)
        self.all_products = normalized
        log.log(f"Normalizados: {len(normalized)} | Duplicados removidos: {duplicates}", "success")

    def _unify(self):
        """Unificación: clasificación bajo taxonomía."""
        for product in self.all_products:
            # Normalizar categoría
            product.categoria = normalize_category(
                product.categoria or product.nombre,
                self.industry
            )

            # Generar descripción corta si falta
            if not product.descripcion_corta and product.descripcion:
                product.descripcion_corta = product.descripcion[:150]

        # Estadísticas de categorías
        categories = {}
        for p in self.all_products:
            categories[p.categoria] = categories.get(p.categoria, 0) + 1

        log.log(f"Categorías: {len(categories)}", "success")
        for cat, count in sorted(categories.items(), key=lambda x: -x[1])[:5]:
            log.log(f"   {cat}: {count}", "debug")

    def _enrich(self):
        """Enriquecimiento: añadir fitment y atributos."""
        for product in self.all_products:
            # Generar ID único
            if not product.id:
                product.id = f"{self.prefix}-{hashlib.md5(f'{product.codigo}{product.nombre}'.encode()).hexdigest()[:8]}"

            # Tags automáticos
            if not product.tags:
                tags = []
                if product.categoria:
                    tags.append(product.categoria.lower())
                if product.marca:
                    tags.append(product.marca.lower())
                words = product.nombre.lower().split()[:3]
                tags.extend(words)
                product.tags = list(set(tags))

        log.log(f"Productos enriquecidos: {len(self.all_products)}", "success")

    def _export(self) -> Tuple[str, str]:
        """Exporta a CSV y JSON."""
        if not self.all_products:
            log.log("No hay productos para exportar", "warning")
            return "", ""

        # Crear DataFrame
        records = [p.to_dict() for p in self.all_products]
        df = pd.DataFrame(records)

        # Ordenar columnas
        priority_cols = ['sku', 'codigo', 'nombre', 'descripcion', 'precio',
                        'categoria', 'marca', 'imagen', 'stock']
        cols = [c for c in priority_cols if c in df.columns]
        cols += [c for c in df.columns if c not in cols]
        df = df[cols]

        # Guardar CSV
        csv_path = os.path.join(self.output_dir, f"{self.prefix}_catalogo.csv")
        df.to_csv(csv_path, sep=';', index=False, encoding='utf-8')

        # Guardar JSON
        json_path = os.path.join(self.output_dir, f"{self.prefix}_catalogo.json")

        # JSON con estructura para Shopify/E-commerce
        export_data = {
            "metadata": {
                "version": VERSION,
                "generated": datetime.now().isoformat(),
                "industry": self.industry,
                "total_products": len(self.all_products),
                "prefix": self.prefix
            },
            "products": records
        }

        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, ensure_ascii=False, indent=2)

        log.log(f"CSV: {csv_path}", "success")
        log.log(f"JSON: {json_path}", "success")

        return csv_path, json_path

    def _print_banner(self):
        print(f"""
{Colors.BOLD}╔══════════════════════════════════════════════════════════════════════╗
║                    🚀 {SCRIPT_NAME} v{VERSION}                       ║
║                 Procesador Universal Multi-Industria                 ║
╚══════════════════════════════════════════════════════════════════════╝{Colors.RESET}
""")
        log.log(f"Industria: {INDUSTRY_TAXONOMIES.get(self.industry, {}).get('name', self.industry)}")
        log.log(f"Prefijo: {self.prefix}")
        log.log(f"Output: {self.output_dir}")

    def _print_summary(self, csv_path: str, json_path: str):
        with_price = sum(1 for p in self.all_products if p.precio > 0)
        with_image = sum(1 for p in self.all_products if p.imagen)

        print(f"""
{Colors.BOLD}{'='*70}
📊 RESUMEN FINAL - FICHA 360°
{'='*70}{Colors.RESET}

{Colors.GREEN}✓ Total productos:{Colors.RESET}     {len(self.all_products)}
{Colors.GREEN}✓ Con precio:{Colors.RESET}          {with_price}
{Colors.GREEN}✓ Con imagen:{Colors.RESET}          {with_image}
{Colors.CYAN}○ Tiempo:{Colors.RESET}              {log.elapsed()}

{Colors.BOLD}📁 ARCHIVOS:{Colors.RESET}
   {csv_path}
   {json_path}

{Colors.BOLD}{'='*70}{Colors.RESET}
""")


# ============================================================================
# CLI
# ============================================================================

def parse_args():
    """Parsea argumentos de línea de comandos."""
    if len(sys.argv) < 2:
        return None

    config = {
        'source': sys.argv[1],
        'output_dir': DEFAULT_OUTPUT_DIR,
        'prefix': 'SRM',
        'industry': 'generico',
        'dpi': 150,
        'pages': None,
    }

    i = 2
    while i < len(sys.argv):
        arg = sys.argv[i]

        if arg in ['--output', '-o'] and i + 1 < len(sys.argv):
            config['output_dir'] = sys.argv[i + 1]
            i += 2
        elif arg == '--prefix' and i + 1 < len(sys.argv):
            config['prefix'] = sys.argv[i + 1].upper()
            i += 2
        elif arg == '--industry' and i + 1 < len(sys.argv):
            config['industry'] = sys.argv[i + 1].lower()
            i += 2
        elif arg == '--dpi' and i + 1 < len(sys.argv):
            config['dpi'] = int(sys.argv[i + 1])
            i += 2
        elif arg == '--pages' and i + 1 < len(sys.argv):
            pages_str = sys.argv[i + 1]
            if pages_str.lower() == 'all':
                config['pages'] = None
            else:
                pages = set()
                for part in pages_str.split(','):
                    if '-' in part:
                        start, end = map(int, part.split('-'))
                        pages.update(range(start, end + 1))
                    else:
                        pages.add(int(part))
                config['pages'] = sorted(pages)
            i += 2
        elif arg == '--scrape':
            config['scrape_pagination'] = True
            i += 1
        else:
            i += 1

    return config


def print_help():
    industries = ", ".join(INDUSTRY_TAXONOMIES.keys())

    print(f"""
{Colors.BOLD}{SCRIPT_NAME} v{VERSION}{Colors.RESET}
{'='*70}

{Colors.CYAN}USO:{Colors.RESET}
    python3 srm_intelligent_processor.py <archivo_o_url> [opciones]

{Colors.CYAN}FORMATOS SOPORTADOS:{Colors.RESET}
    PDF, Excel (.xlsx/.xls), CSV, Word (.docx), TXT,
    Imágenes (JPG/PNG/WEBP), ZIP, URLs

{Colors.CYAN}OPCIONES:{Colors.RESET}
    --output, -o DIR      Directorio de salida
    --prefix PREFIX       Prefijo para SKU (default: SRM)
    --industry INDUSTRY   Industria: {industries}
    --dpi DPI             Resolución para PDFs (default: 150)
    --pages PAGES         Páginas de PDF: "2-50", "1,3,5", "all"
    --scrape              Seguir paginación en URLs
    --help, -h            Mostrar esta ayuda

{Colors.CYAN}EJEMPLOS:{Colors.RESET}
    # Procesar PDF de autopartes
    python3 srm_intelligent_processor.py catalogo.pdf --industry autopartes --prefix ARM

    # Procesar Excel de ferretería
    python3 srm_intelligent_processor.py productos.xlsx --industry ferreteria

    # Scraping de competencia
    python3 srm_intelligent_processor.py "https://competencia.com/catalogo" --scrape

    # Procesar ZIP con imágenes
    python3 srm_intelligent_processor.py fotos_productos.zip --industry electronica

{Colors.CYAN}VARIABLES DE ENTORNO:{Colors.RESET}
    OPENAI_API_KEY       API key de OpenAI (requerido)

{Colors.CYAN}DEPENDENCIAS:{Colors.RESET}
    pip install openai pandas openpyxl python-docx beautifulsoup4 requests pillow opencv-python numpy
    apt install poppler-utils  # Para PDFs
""")


def main():
    if '--help' in sys.argv or '-h' in sys.argv or len(sys.argv) < 2:
        print_help()
        sys.exit(0)

    config = parse_args()
    if not config:
        print_help()
        sys.exit(1)

    # Validar API key
    if not os.getenv("OPENAI_API_KEY"):
        log.log("OPENAI_API_KEY no configurada", "error")
        sys.exit(1)

    # Ejecutar pipeline
    try:
        pipeline = SRMPipeline(config)
        pipeline.process(config['source'])
    except KeyboardInterrupt:
        log.log("\nProceso interrumpido", "warning")
        sys.exit(130)
    except Exception as e:
        log.log(f"Error: {e}", "error")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
